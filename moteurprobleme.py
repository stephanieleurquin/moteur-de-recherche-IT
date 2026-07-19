import streamlit as st
import sqlite3
import pandas as pd
import re
from datetime import datetime, date
import hashlib
import random
import string
from io import BytesIO
import textwrap
import os

# ==================================================
# CONFIGURATION
# ==================================================

st.set_page_config(
    page_title="Assistant IT Pro - Premium",
    page_icon="💻",
    layout="wide"
)

DB = "assistant_it_pro.db"

# ==================================================
# FONCTIONS DE BASE DE DONNÉES (DOIT ÊTRE ICI)
# ==================================================

def connexion_db():
    try:
        conn = sqlite3.connect(DB)
        return conn
    except sqlite3.Error as e:
        st.error(f"Erreur de connexion : {e}")
        return None

# ==================================================
# CSS
# ==================================================

st.markdown("""
<style>
    .stApp { background-color: #0a0a0f; }
    h1, h2, h3 { color: #00d4ff !important; }
    p, li, label { color: #ffffff !important; }

    .stButton > button {
        background: linear-gradient(135deg, #00d4ff 0%, #0077be 100%) !important;
        color: #0a0a0f !important;
        font-weight: 700 !important;
        border: none !important;
        border-radius: 50px !important;
        padding: 12px 30px !important;
    }
    .stButton > button:hover {
        transform: scale(1.05) !important;
        box-shadow: 0 0 30px #00d4ff40 !important;
    }

    .stTextInput > div > div > input, .stTextArea > div > div > textarea {
        background: #1a1a2e !important;
        border: 2px solid #2a2a4a !important;
        border-radius: 12px !important;
        color: white !important;
    }
    .stTextInput > div > div > input:focus, .stTextArea > div > div > textarea:focus {
        border-color: #00d4ff !important;
    }
</style>
""", unsafe_allow_html=True)

# ==================================================
# SESSION STATE
# ==================================================

if "user" not in st.session_state:
    st.session_state.user = None
if "premium" not in st.session_state:
    st.session_state.premium = False
if "plan" not in st.session_state:
    st.session_state.plan = "gratuit"
if "recherches" not in st.session_state:
    st.session_state.recherches = 0
if "recherches_jour" not in st.session_state:
    st.session_state.recherches_jour = 0
if "date_recherche" not in st.session_state:
    st.session_state.date_recherche = date.today()
if "page" not in st.session_state:
    st.session_state.page = "accueil"
if "moteur" not in st.session_state:
    st.session_state.moteur = None
if "montant_virement" not in st.session_state:
    st.session_state.montant_virement = 0
if "offre_virement" not in st.session_state:
    st.session_state.offre_virement = ""
if "plan_virement" not in st.session_state:
    st.session_state.plan_virement = ""
if "ref_virement" not in st.session_state:
    st.session_state.ref_virement = ""

# ==================================================
# OFFRES
# ==================================================

OFFRES = {
    "gratuit": {
        "nom": "Gratuit",
        "prix": "0€",
        "recherches": 3,
        "features": ["3 recherches par jour", "70+ diagnostics", "Diagnostics basiques"]
    },
    "pro": {
        "nom": "Pro",
        "prix": "9.90€/mois",
        "recherches": 999,
        "features": ["Recherches illimitées", "150+ diagnostics", "Diagnostics avancés", "Export PDF",
                     "Support prioritaire"]
    },
    "business": {
        "nom": "Business",
        "prix": "29.90€/mois",
        "recherches": 9999,
        "features": ["Recherches illimitées", "150+ diagnostics", "Diagnostics experts", "Export PDF/Word",
                     "Support 24/7", "Accès API", "5 comptes inclus"]
    }
}

# ==================================================
# CRÉATION ET REMPLISSAGE DE LA BASE
# ==================================================

def creer_base():
    conn = connexion_db()
    if conn is None:
        return
    cur = conn.cursor()
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS pannes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            titre TEXT,
            description TEXT,
            diagnostic TEXT,
            procedure TEXT,
            questions TEXT,
            categorie TEXT,
            niveau INTEGER,
            tags TEXT
        )
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS entreprises (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT,
            date_creation TEXT,
            plan TEXT DEFAULT 'business'
        )
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS utilisateurs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE,
            password TEXT,
            plan TEXT DEFAULT 'gratuit',
            premium INTEGER DEFAULT 0,
            recherches INTEGER DEFAULT 0,
            date_inscription TEXT,
            entreprise_id INTEGER,
            role TEXT DEFAULT 'membre',
            abonnement_expire_le TEXT,
            FOREIGN KEY (entreprise_id) REFERENCES entreprises (id)
        )
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS invitations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT,
            token TEXT UNIQUE,
            entreprise_id INTEGER,
            date_creation TEXT,
            expire_le TEXT,
            FOREIGN KEY (entreprise_id) REFERENCES entreprises (id)
        )
    """)
    
    conn.commit()
    try:
        cur.execute("ALTER TABLE utilisateurs ADD COLUMN abonnement_expire_le TEXT")
    except:
        pass
    conn.close()

def remplir_base():
    conn = connexion_db()
    if conn is None:
        return
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM pannes")
    if cur.fetchone()[0] == 0:
        donnees = [
            ("Windows ne s'installe pas", "L'installation de Windows échoue", "Problème de clé USB ou de pilote manquant", "1- Vérifier la clé USB\n2- Désactiver Secure Boot\n3- Installer les pilotes manuellement", "Quelle version de Windows ?", "Windows", 4, "installation,windows"),
            ("Windows demande une clé", "Windows demande une clé d'activation", "Clé manquante ou invalide", "1- Vérifier la clé\n2- Contacter Microsoft\n3- Utiliser une clé générique", "Avez-vous une clé valide ?", "Windows", 2, "cle,activation,windows"),
            ("Windows Update échoue", "Les mises à jour Windows ne s'installent pas", "Problème de cache ou de service", "1- Arrêter le service Update\n2- Vider le cache\n3- Redémarrer le service", "Quelle est l'erreur ?", "Windows", 3, "update,windows"),
            # Ajoute ici toutes tes autres pannes (je mets juste un échantillon)
        ]
        cur.executemany(
            "INSERT INTO pannes (titre, description, diagnostic, procedure, questions, categorie, niveau, tags) VALUES (?,?,?,?,?,?,?,?)",
            donnees
        )
        conn.commit()
    conn.close()

# ==================================================
# MOTEUR DE RECHERCHE
# ==================================================

class RechercheIT:
    def __init__(self):
        self.df = None

    def charger(self):
        if self.df is None:
            conn = connexion_db()
            if conn is None:
                return
            self.df = pd.read_sql_query("SELECT * FROM pannes", conn)
            conn.close()

    def rechercher(self, question):
        self.charger()
        if self.df is None or self.df.empty:
            return []
        question = question.lower()
        mots = re.findall(r"\w+", question)
        resultats = []
        for _, panne in self.df.iterrows():
            score = 0
            champs = f"{panne['titre']} {panne['description']} {panne['tags']} {panne['categorie']}".lower()
            for mot in mots:
                if len(mot) > 1 and mot in champs:
                    score += 5
                if mot in panne['titre'].lower():
                    score += 10
            if score > 0:
                resultats.append((dict(panne), score))
        resultats.sort(key=lambda x: x[1], reverse=True)
        return resultats[:10]

# ==================================================
# EXPORT PDF / WORD
# ==================================================

def generer_pdf_resultats(resultats, question):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import simpleSplit
    except ImportError:
        st.error("❌ reportlab non installé")
        return None
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    largeur, hauteur = A4
    y = hauteur - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Résultats de recherche - Assistant IT Pro")
    y -= 30
    c.setFont("Helvetica", 12)
    c.drawString(50, y, f"Question : {question}")
    y -= 20
    c.drawString(50, y, f"{len(resultats)} résultat(s)")
    y -= 30
    for i, (panne, score) in enumerate(resultats, 1):
        if y < 100:
            c.showPage()
            y = hauteur - 50
            c.setFont("Helvetica", 12)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, f"{i}. {panne['titre']} (Score: {score})")
        y -= 20
        c.setFont("Helvetica", 10)
        for ligne in simpleSplit(f"Catégorie : {panne['categorie']}", "Helvetica", 10, largeur-100):
            c.drawString(60, y, ligne); y -= 15
        for ligne in simpleSplit(f"Diagnostic : {panne['diagnostic']}", "Helvetica", 10, largeur-100):
            c.drawString(60, y, ligne); y -= 15
        for ligne in simpleSplit(f"Procédure : {panne['procedure']}", "Helvetica", 10, largeur-100):
            c.drawString(60, y, ligne); y -= 15
        if panne.get('questions'):
            for ligne in simpleSplit(f"Questions : {panne['questions']}", "Helvetica", 10, largeur-100):
                c.drawString(60, y, ligne); y -= 15
        y -= 10
    c.save()
    return buffer.getvalue()

def generer_word_resultats(resultats, question):
    try:
        import docx
    except ImportError:
        st.error("❌ python-docx non installé")
        return None
    doc = docx.Document()
    doc.add_heading("Résultats de recherche - Assistant IT Pro", 0)
    doc.add_paragraph(f"Question : {question}")
    doc.add_paragraph(f"{len(resultats)} résultat(s)")
    for i, (panne, score) in enumerate(resultats, 1):
        doc.add_heading(f"{i}. {panne['titre']} (Score: {score})", level=1)
        p = doc.add_paragraph(); p.add_run("Catégorie : ").bold = True; p.add_run(panne['categorie'])
        p = doc.add_paragraph(); p.add_run("Diagnostic : ").bold = True; p.add_run(panne['diagnostic'])
        p = doc.add_paragraph(); p.add_run("Procédure : ").bold = True; p.add_run(panne['procedure'])
        if panne.get('questions'):
            p = doc.add_paragraph(); p.add_run("Questions : ").bold = True; p.add_run(panne['questions'])
        doc.add_paragraph()
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ==================================================
# AUTHENTIFICATION
# ==================================================

def inscription(email, password):
    try:
        conn = connexion_db()
        if conn is None:
            return False
        cur = conn.cursor()
        pwd = hashlib.sha256(password.encode()).hexdigest()
        cur.execute(
            "INSERT INTO utilisateurs (email, password, plan, premium, recherches, date_inscription) VALUES (?, ?, 'gratuit', 0, 0, ?)",
            (email, pwd, date.today().isoformat())
        )
        conn.commit()
        conn.close()
        return True
    except sqlite3.IntegrityError:
        return False
    except Exception as e:
        st.error(f"Erreur inscription : {e}")
        return False

def connexion_utilisateur(email, password):
    conn = connexion_db()
    if conn is None:
        return None
    cur = conn.cursor()
    pwd = hashlib.sha256(password.encode()).hexdigest()
    cur.execute("SELECT * FROM utilisateurs WHERE email = ? AND password = ?", (email, pwd))
    user = cur.fetchone()
    if user:
        expire_le = user[9] if len(user) > 9 else None
        if expire_le:
            try:
                date_expiration = datetime.fromisoformat(expire_le)
                if datetime.now() > date_expiration:
                    cur.execute("UPDATE utilisateurs SET plan='gratuit', premium=0, abonnement_expire_le=NULL WHERE email=?", (email,))
                    conn.commit()
                    cur.execute("SELECT * FROM utilisateurs WHERE email = ? AND password = ?", (email, pwd))
                    user = cur.fetchone()
            except:
                pass
    conn.close()
    return user

def mise_a_jour_plan(email, plan):
    conn = connexion_db()
    if conn is None:
        return
    cur = conn.cursor()
    cur.execute("UPDATE utilisateurs SET plan = ?, premium = 1 WHERE email = ?", (plan, email))
    if plan == "business":
        cur.execute("SELECT entreprise_id FROM utilisateurs WHERE email = ?", (email,))
        result = cur.fetchone()
        if result and result[0] is None:
            nom_entreprise = f"Entreprise de {email}"
            cur.execute("INSERT INTO entreprises (nom, date_creation, plan) VALUES (?, ?, 'business')",
                        (nom_entreprise, date.today().isoformat()))
            entreprise_id = cur.lastrowid
            cur.execute("UPDATE utilisateurs SET entreprise_id = ?, role = 'admin' WHERE email = ?",
                        (entreprise_id, email))
    conn.commit()
    conn.close()

# ==================================================
# PAGES
# ==================================================

def page_virement():
    st.markdown('<p style="color:#FFD700; font-size:36px; font-weight:700; text-align:center;">💳 Paiement par Virement Bancaire</p>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("""
    <div style='background: linear-gradient(135deg, #1a5276 0%, #2e86c1 100%); padding: 30px; border-radius: 20px; text-align: center; margin-bottom: 30px;'>
        <h2 style='color: white;'>Paiement sécurisé</h2>
        <p style='color: #FFD700;'>Virement bancaire - 0€ de frais</p>
    </div>
    """, unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        with st.container(border=True):
            st.markdown("### 🚀 PRO\n**9.90€ / mois** | **79€ / an**\n- Recherches illimitées\n- Diagnostics avancés\n- 150+ diagnostics\n- Export PDF\n- Support prioritaire\n- Statistiques avancées")
            if st.button("Choisir Pro - 9.90€", type="primary", use_container_width=True):
                st.session_state.montant_virement = 9.90
                st.session_state.offre_virement = "Pro"
                st.session_state.plan_virement = "pro"
                st.success("✅ Offre Pro sélectionnée !")
                st.balloons()
    with col2:
        with st.container(border=True):
            st.markdown("### 🏢 BUSINESS\n**29.90€ / mois** | **249€ / an**\n- Tout Pro inclus\n- Diagnostics experts\n- 150+ diagnostics\n- Export PDF/Word\n- Support 24/7\n- Accès API\n- 5 comptes inclus")
            if st.button("Choisir Business - 29.90€", type="primary", use_container_width=True):
                st.session_state.montant_virement = 29.90
                st.session_state.offre_virement = "Business"
                st.session_state.plan_virement = "business"
                st.success("✅ Offre Business sélectionnée !")
                st.balloons()
    if st.session_state.montant_virement > 0:
        st.markdown("---")
        ref = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
        st.session_state.ref_virement = ref
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("""
            <div style='background: #1a1a2e; padding: 20px; border-radius: 10px; border: 1px solid #2a2a4a;'>
                <h4 style='color: #00d4ff;'>Coordonnées bancaires</h4>
                <p><strong style='color:#aaa;'>Titulaire :</strong> <span style='color:white;'>IT Pro Solutions</span></p>
                <p><strong style='color:#aaa;'>IBAN :</strong> <span style='color:white;'>BE80 9733 8252 3877</span></p>
                <p><strong style='color:#aaa;'>BIC :</strong> <span style='color:white;'>ARSPBE22XXX</span></p>
                <p><strong style='color:#aaa;'>Banque :</strong> <span style='color:white;'>ARGENTA</span></p>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown(f"""
            <div style='background: #1a1a2e; padding: 20px; border-radius: 10px; border: 1px solid #2ecc71;'>
                <h4 style='color: #2ecc71;'>Informations</h4>
                <p><strong style='color:#aaa;'>Montant :</strong> <span style='color:#00d4ff;font-weight:700;'>{st.session_state.montant_virement}€</span></p>
                <p><strong style='color:#aaa;'>Offre :</strong> <span style='color:#FFD700;'>{st.session_state.offre_virement}</span></p>
                <p><strong style='color:#aaa;'>Référence :</strong> <code style='background:#0a0a0f;color:#00d4ff;padding:2px 8px;border-radius:4px;'>{ref}</code></p>
                <p><strong style='color:#aaa;'>Email :</strong> <span style='color:white;'>tech.contactinformatique@proton.me</span></p>
                <p style='color: #e74c3c; font-weight:700;'>⚠️ Indiquez la référence</p>
            </div>
            """, unsafe_allow_html=True)
        st.info(f"**Résumé :** {st.session_state.montant_virement}€ - {st.session_state.offre_virement} - Réf: {ref}")
        st.warning("⏳ Activation sous 24-48h ouvrés.")

def page_offres():
    st.markdown('<p style="color:#FFD700; font-size:36px; font-weight:700; text-align:center;">📋 Nos Offres</p>', unsafe_allow_html=True)
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("""
        <div style='background: #1a1a2e; padding: 20px; border-radius: 15px; border: 1px solid #2a2a4a; text-align: center;'>
            <h3 style='color: #aaa;'>🆓 Gratuit</h3>
            <p style='font-size: 28px; color: white;'>0€</p>
            <hr>
            <p style='color: #ccc;'>✅ 3 recherches / jour</p>
            <p style='color: #ccc;'>✅ 70+ diagnostics</p>
            <p style='color: #ccc;'>✅ Diagnostics basiques</p>
            <br>
            <span style='background: #2a2a4a; padding: 8px 20px; border-radius: 50px; color: #aaa;'>Actuel</span>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown("""
        <div style='background: #1a1a2e; padding: 20px; border-radius: 15px; border: 2px solid #FFD700; text-align: center;'>
            <h3 style='color: #FFD700;'>🚀 Pro</h3>
            <p style='font-size: 28px; color: white;'>9.90€<span style='font-size: 16px; color: #aaa;'> /mois</span></p>
            <hr>
            <p style='color: #ccc;'>✅ Recherches illimitées</p>
            <p style='color: #ccc;'>✅ 150+ diagnostics</p>
            <p style='color: #ccc;'>✅ Diagnostics avancés</p>
            <p style='color: #ccc;'>✅ Export PDF</p>
            <p style='color: #ccc;'>✅ Support prioritaire</p>
            <br>
        """, unsafe_allow_html=True)
        if st.button("Choisir Pro", type="primary", key="offre_pro"):
            st.session_state.montant_virement = 9.90
            st.session_state.offre_virement = "Pro"
            st.session_state.plan_virement = "pro"
            st.session_state.page = "💳 Virement"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
    with col3:
        st.markdown("""
        <div style='background: #1a1a2e; padding: 20px; border-radius: 15px; border: 2px solid #9B59B6; text-align: center;'>
            <h3 style='color: #9B59B6;'>🏢 Business</h3>
            <p style='font-size: 28px; color: white;'>29.90€<span style='font-size: 16px; color: #aaa;'> /mois</span></p>
            <hr>
            <p style='color: #ccc;'>✅ Tout Pro inclus</p>
            <p style='color: #ccc;'>✅ Diagnostics experts</p>
            <p style='color: #ccc;'>✅ Export PDF/Word</p>
            <p style='color: #ccc;'>✅ Support 24/7</p>
            <p style='color: #ccc;'>✅ Accès API</p>
            <p style='color: #ccc;'>✅ 5 comptes inclus</p>
            <br>
        """, unsafe_allow_html=True)
        if st.button("Choisir Business", type="primary", key="offre_business"):
            st.session_state.montant_virement = 29.90
            st.session_state.offre_virement = "Business"
            st.session_state.plan_virement = "business"
            st.session_state.page = "💳 Virement"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
    st.info("💡 Offres sans engagement.")

def page_licence():
    st.markdown('<p style="color:#FFD700; font-size:36px; font-weight:700; text-align:center;">📄 Licence et Mentions légales</p>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("""
    ### 📌 Propriété intellectuelle
    - Tous droits réservés à **IT Pro Solutions**.
    - Reproduction interdite.

    ### 🔒 Protection des données
    - Conforme RGPD.

    ### 💰 Paiements
    - Par virement bancaire.

    ### 📞 Support
    - tech.contactinformatique@proton.me

    ---
    *Version 2.0 – 2026*
    """)

# ==================================================
# MAIN
# ==================================================

def main():
    creer_base()
    remplir_base()

    if st.session_state.moteur is None:
        st.session_state.moteur = RechercheIT()

    if st.session_state.date_recherche != date.today():
        st.session_state.date_recherche = date.today()
        st.session_state.recherches_jour = 0

    # ===== SIDEBAR =====
    with st.sidebar:
        st.markdown("""
        <style> section[data-testid="stSidebar"] { background-color: #1a1a2e !important; } </style>
        """, unsafe_allow_html=True)
        st.markdown('<p style="color:#1458; font-size:24px; font-weight:700; text-align:center;">💻 IT Pro</p>', unsafe_allow_html=True)
        st.markdown('<p style="color:#AAAAAA; font-size:12px; text-align:center;">1000 diagnostics</p>', unsafe_allow_html=True)
        st.markdown("---")

        if st.session_state.user:
            st.markdown(f'<p style="color:#FFFFFF;">👤 {st.session_state.user}</p>', unsafe_allow_html=True)
            plan = st.session_state.plan
            if plan == "business":
                st.markdown('<div style="background:#9B59B6; padding:12px; border-radius:10px; text-align:center;"><p style="color:white; font-weight:700; margin:0;">🏢 BUSINESS</p></div>', unsafe_allow_html=True)
            elif plan == "pro":
                st.markdown('<div style="background:#FFD700; padding:12px; border-radius:10px; text-align:center;"><p style="color:#0a0a0f; font-weight:700; margin:0;">🚀 PRO</p></div>', unsafe_allow_html=True)
            else:
                st.markdown('<div style="background:#FF6B6B; padding:12px; border-radius:10px; text-align:center;"><p style="color:white; font-weight:700; margin:0;">🆓 GRATUIT</p></div>', unsafe_allow_html=True)
                restant = max(0, 3 - st.session_state.recherches_jour)
                st.markdown(f'<p style="color:#FFFFFF;">🔍 {restant} recherches restantes</p>', unsafe_allow_html=True)
                st.progress(st.session_state.recherches_jour / 3)
            st.markdown("---")
            menu = ["🏠 Accueil", "📋 Offres", "💳 Virement", "📄 Licence"]
            st.session_state.page = st.radio("Navigation", menu, key="sidebar_menu")
            if st.button("🚪 Déconnexion", use_container_width=True):
                st.session_state.user = None
                st.session_state.premium = False
                st.session_state.plan = "gratuit"
                st.session_state.recherches = 0
                st.session_state.recherches_jour = 0
                st.rerun()
        else:
            tab1, tab2 = st.tabs(["🔐 Connexion", "📝 Inscription"])
            with tab1:
                email = st.text_input("Email", key="login_email")
                password = st.text_input("Mot de passe", type="password", key="login_pass")
                if st.button("Se connecter", use_container_width=True):
                    user = connexion_utilisateur(email, password)
                    if user:
                        st.session_state.user = email
                        st.session_state.plan = user[3] if user[3] else "gratuit"
                        st.session_state.premium = bool(user[4])
                        st.session_state.recherches = user[5] if user[5] else 0
                        st.session_state.recherches_jour = user[5] if user[5] else 0
                        st.success("✅ Connecté !")
                        st.rerun()
                    else:
                        st.error("❌ Identifiants incorrects")
            with tab2:
                email = st.text_input("Email", key="register_email")
                password = st.text_input("Mot de passe", type="password", key="register_pass")
                if st.button("Créer un compte", use_container_width=True):
                    if inscription(email, password):
                        st.success("✅ Compte créé ! Connectez-vous")
                    else:
                        st.error("❌ Email déjà utilisé")

    # ===== GESTION DES PAGES =====
    if st.session_state.page == "📋 Offres":
        page_offres()
        return
    if st.session_state.page == "💳 Virement":
        page_virement()
        return
    if st.session_state.page == "📄 Licence":
        page_licence()
        return

    # ===== ACCUEIL =====
    st.markdown('<p style="color:#00d4ff; font-size:48px; font-weight:900; text-align:center;">🔧 Assistant Dépannage IT</p>', unsafe_allow_html=True)
    st.markdown('<p style="color:#aaa; text-align:center; font-size:18px;">Par IT Pro Solutions - <span style="color:#FFD700;">150+ diagnostics</span></p>', unsafe_allow_html=True)
    st.markdown("---")

    question = st.text_area("Décrivez votre problème :", height=100, placeholder="Ex: mon PC est lent, le wifi ne marche pas, erreur Windows...")

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("🔍 Rechercher", type="primary", use_container_width=True):
            if not st.session_state.user:
                st.error("❌ Connectez-vous d'abord")
            elif not st.session_state.premium and st.session_state.recherches_jour >= 3:
                st.error("🔴 LIMITE ATTEINTE ! Passez Premium pour continuer.")
                if st.button("VOIR LES OFFRES"):
                    st.session_state.page = "📋 Offres"
                    st.rerun()
            elif question.strip():
                with st.spinner("Recherche en cours..."):
                    if not st.session_state.premium:
                        st.session_state.recherches_jour += 1
                        st.session_state.recherches += 1
                    results = st.session_state.moteur.rechercher(question)
                    if results:
                        st.success(f"✅ {len(results)} résultat(s) trouvé(s)")
                        for panne, score in results:
                            with st.expander(f"🔹 {panne['titre']} (Score: {score})"):
                                st.markdown(f"**Catégorie:** {panne['categorie']}")
                                st.markdown(f"**Diagnostic:** {panne['diagnostic']}")
                                st.markdown(f"**Procédure:**\n{panne['procedure']}")
                                if panne.get('questions'):
                                    st.info(f"❓ {panne['questions']}")
                        # Exports
                        if st.session_state.plan in ["pro", "business"]:
                            st.markdown("---")
                            col_btn1, col_btn2 = st.columns(2)
                            with col_btn1:
                                pdf_data = generer_pdf_resultats(results, question)
                                if pdf_data:
                                    st.download_button("📄 Télécharger en PDF", pdf_data, f"resultats_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", "application/pdf", key="pdf_download")
                                else:
                                    st.warning("Export PDF indisponible")
                            with col_btn2:
                                word_data = generer_word_resultats(results, question)
                                if word_data:
                                    st.download_button("📝 Télécharger en Word", word_data, f"resultats_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="word_download")
                                else:
                                    st.warning("Export Word indisponible")
                        else:
                            st.info("🔒 Export PDF/Word pour abonnés **Pro** et **Business**.")
                    else:
                        st.warning("😕 Aucun résultat trouvé")

    st.markdown("---")
    st.markdown('<p style="text-align:center; color:#444; font-size:12px;">© 2026 <strong style="color:#FFD700;">IT Pro Solutions</strong> - Tous droits réservés</p>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
